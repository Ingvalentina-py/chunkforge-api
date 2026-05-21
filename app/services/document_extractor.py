import io
from dataclasses import dataclass

from docx import Document
from fastapi import HTTPException, UploadFile
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


@dataclass
class ExtractedDocument:
    text: str
    filename: str
    page_texts: list[str] | None = None


def _decode_text_bytes(data: bytes) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("latin-1")


def _extract_pdf(data: bytes) -> tuple[str, list[str]]:
    reader = PdfReader(io.BytesIO(data))
    page_texts: list[str] = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        page_texts.append(page_text.strip())
    full_text = "\n\n".join(t for t in page_texts if t)
    return full_text, page_texts


def _extract_docx(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


async def extract_from_upload(file: UploadFile) -> ExtractedDocument:
    filename = file.filename or "document"
    extension = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if extension not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Formato no soportado. Use: {', '.join(sorted(SUPPORTED_EXTENSIONS))}",
        )

    data = await file.read()
    if not data:
        raise HTTPException(status_code=422, detail="El archivo está vacío")

    if extension == ".pdf":
        text, page_texts = _extract_pdf(data)
        if not text.strip():
            raise HTTPException(
                status_code=422,
                detail="No se pudo extraer texto del PDF",
            )
        return ExtractedDocument(text=text, filename=filename, page_texts=page_texts)

    if extension == ".docx":
        text = _extract_docx(data)
    else:
        text = _decode_text_bytes(data)

    if not text.strip():
        raise HTTPException(
            status_code=422,
            detail="No se pudo extraer texto del documento",
        )

    return ExtractedDocument(text=text, filename=filename, page_texts=None)


def extract_from_text(text: str, filename: str = "inline.txt") -> ExtractedDocument:
    if not text.strip():
        raise HTTPException(status_code=422, detail="El texto proporcionado está vacío")
    return ExtractedDocument(text=text, filename=filename, page_texts=None)
